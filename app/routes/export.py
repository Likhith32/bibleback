"""
Export routes — PDF and Word document generation.
"""
import io
import logging
from fastapi import APIRouter, Body, HTTPException, Response
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdf_canvas
from docx import Document

logger = logging.getLogger(__name__)
router = APIRouter()

PAGE_W, PAGE_H = A4
LEFT_MARGIN = 25 * mm
TOP_MARGIN = PAGE_H - 25 * mm
LINE_HEIGHT = 14


@router.post("/export/pdf")
async def export_pdf(payload: dict = Body(...)):
    content = payload.get("content", "")
    title = payload.get("title", "Bible Study Notes")

    if not content:
        raise HTTPException(status_code=400, detail="content is required")

    try:
        buf = io.BytesIO()
        c = pdf_canvas.Canvas(buf, pagesize=A4)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(LEFT_MARGIN, TOP_MARGIN, title)
        c.setFont("Helvetica", 10)

        y = TOP_MARGIN - 30
        for line in content.split("\n"):
            # Word-wrap long lines
            while len(line) > 90:
                c.drawString(LEFT_MARGIN, y, line[:90])
                line = line[90:]
                y -= LINE_HEIGHT
                if y < 40:
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = TOP_MARGIN
            c.drawString(LEFT_MARGIN, y, line)
            y -= LINE_HEIGHT
            if y < 40:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = TOP_MARGIN

        c.save()
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{title}.pdf"'},
        )

    except Exception as e:
        logger.error("PDF export error: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="PDF export failed")


@router.post("/export/docx")
async def export_docx(payload: dict = Body(...)):
    content = payload.get("content", "")
    title = payload.get("title", "Bible Study Notes")

    if not content:
        raise HTTPException(status_code=400, detail="content is required")

    try:
        doc = Document()
        doc.add_heading(title, level=0)

        for paragraph_text in content.split("\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
        )

    except Exception as e:
        logger.error("DOCX export error: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail="DOCX export failed")
